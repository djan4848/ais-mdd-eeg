#!/usr/bin/env python3
"""
generate_highlighted_changes.py
Compare AIS_MDD_manuscript_v13_final.docx vs v14.docx and produce
AIS_MDD_manuscript_v14_highlighted_changes.docx with:
  - Yellow highlight  = text added or modified in v14
  - Red strikethrough = text present in v13 that was removed/replaced
"""

import os
import difflib
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = '/media/neuraldyn/PortableSSD/DEPRESSION'
V13_PATH = os.path.join(BASE_DIR, 'AIS_MDD_manuscript_v13_final.docx')
V14_PATH = os.path.join(BASE_DIR, 'AIS_MDD_manuscript_v14.docx')
OUT_PATH = os.path.join(BASE_DIR, 'AIS_MDD_manuscript_v14_highlighted_changes.docx')

SIMILARITY_THRESHOLD = 0.999  # highlight any paragraph with even a single character change


# ── XML helpers ────────────────────────────────────────────────────────────────

def add_highlight_yellow(run):
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    rPr.append(highlight)


def add_strikethrough_red(para, text):
    """Prepend red strikethrough annotation before existing runs."""
    if not text.strip():
        return
    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Strike
    strike = OxmlElement('w:strike')
    rPr.append(strike)
    # Color red
    color = OxmlElement('w:color')
    color.set(qn('w:val'), 'FF0000')
    rPr.append(color)
    # Size 9pt = 18 half-points
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    r_elem.append(rPr)
    t_elem = OxmlElement('w:t')
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = f'[DELETED: {text[:200]}{"…" if len(text) > 200 else ""}] '
    r_elem.append(t_elem)
    # Insert before first child run
    p_elem = para._p
    first_r = p_elem.find(qn('w:r'))
    if first_r is not None:
        p_elem.insert(list(p_elem).index(first_r), r_elem)
    else:
        p_elem.append(r_elem)


def highlight_paragraph(para):
    """Highlight all non-empty runs in yellow."""
    for run in para.runs:
        if run.text.strip():
            add_highlight_yellow(run)
    # Fallback: paragraph has text but no runs (uncommon)
    if not para.runs and para.text.strip():
        run = para.add_run(para.text)
        add_highlight_yellow(run)


def highlight_cell(cell):
    for para in cell.paragraphs:
        highlight_paragraph(para)


# ── Load documents ─────────────────────────────────────────────────────────────

print("Loading documents…")
doc_v13 = Document(V13_PATH)
doc_v14 = Document(V14_PATH)
doc_out = Document(V14_PATH)   # working copy to annotate

# ── Paragraph-level diff ───────────────────────────────────────────────────────

paras_v13 = [p.text.strip() for p in doc_v13.paragraphs]
paras_v14 = [p.text.strip() for p in doc_v14.paragraphs]

matcher = difflib.SequenceMatcher(None, paras_v13, paras_v14, autojunk=False)
opcodes = matcher.get_opcodes()

changed_indices_v14 = set()   # paragraph indices in v14 to highlight
deleted_before = {}            # j1 → deleted text from v13

for tag, i1, i2, j1, j2 in opcodes:
    if tag == 'replace':
        for j in range(j1, j2):
            old = paras_v13[i1] if i1 < len(paras_v13) else ''
            new = paras_v14[j]   if j  < len(paras_v14)  else ''
            ratio = difflib.SequenceMatcher(None, old, new).ratio()
            if ratio < SIMILARITY_THRESHOLD:
                changed_indices_v14.add(j)
        # Collect deleted text (first replaced paragraph only)
        deleted_text = ' / '.join(
            paras_v13[i] for i in range(i1, i2) if paras_v13[i].strip()
        )
        if deleted_text and j1 not in deleted_before:
            deleted_before[j1] = deleted_text

    elif tag == 'insert':
        for j in range(j1, j2):
            changed_indices_v14.add(j)

print(f"Paragraphs modified/inserted in v14: {len(changed_indices_v14)}")

# ── Apply paragraph highlights ─────────────────────────────────────────────────

paras_out = doc_out.paragraphs
changes_applied = 0

for idx, para in enumerate(paras_out):
    if idx in changed_indices_v14 and para.text.strip():
        if idx in deleted_before:
            add_strikethrough_red(para, deleted_before[idx])
        highlight_paragraph(para)
        changes_applied += 1

# ── Table cell diff ────────────────────────────────────────────────────────────

tables_v13 = doc_v13.tables
tables_v14 = doc_v14.tables
tables_out = doc_out.tables
table_changes = 0

for t_idx in range(min(len(tables_v13), len(tables_v14), len(tables_out))):
    tbl13, tbl14, tbl_out = tables_v13[t_idx], tables_v14[t_idx], tables_out[t_idx]
    try:
        for r_idx in range(min(len(tbl13.rows), len(tbl14.rows), len(tbl_out.rows))):
            row13, row14, row_out = tbl13.rows[r_idx], tbl14.rows[r_idx], tbl_out.rows[r_idx]
            for c_idx in range(min(len(row13.cells), len(row14.cells), len(row_out.cells))):
                c13 = row13.cells[c_idx].text.strip()
                c14 = row14.cells[c_idx].text.strip()
                if c13 != c14:
                    ratio = difflib.SequenceMatcher(None, c13, c14).ratio()
                    if ratio < SIMILARITY_THRESHOLD:
                        highlight_cell(row_out.cells[c_idx])
                        table_changes += 1
    except Exception as e:
        print(f"  Table {t_idx}: {e}")

print(f"Table cells modified: {table_changes}")

# ── Revision tracking note at top ─────────────────────────────────────────────

note_p = OxmlElement('w:p')
note_pPr = OxmlElement('w:pPr')
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'Normal')
note_pPr.append(pStyle)
note_p.append(note_pPr)

note_r = OxmlElement('w:r')
note_rPr = OxmlElement('w:rPr')
bold = OxmlElement('w:b')
note_rPr.append(bold)
color = OxmlElement('w:color')
color.set(qn('w:val'), '7B2D00')
note_rPr.append(color)
# Yellow background on note itself
hl = OxmlElement('w:highlight')
hl.set(qn('w:val'), 'yellow')
note_rPr.append(hl)
note_r.append(note_rPr)
note_t = OxmlElement('w:t')
note_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
note_t.text = (
    "REVISION TRACKING NOTE — Yellow highlight = text added or modified in v14. "
    "Red strikethrough [DELETED: …] = text present in v13 that was removed or replaced. "
    f"Threshold: changes <{int(SIMILARITY_THRESHOLD*100)}% similarity are marked. "
    "Manuscript: JADR-D-26-00295 | v13→v14 comparison."
)
note_r.append(note_t)
note_p.append(note_r)
doc_out.element.body.insert(0, note_p)

# ── Save ───────────────────────────────────────────────────────────────────────

doc_out.save(OUT_PATH)
print(f"\nSaved: {OUT_PATH}")
print(f"Paragraphs highlighted: {changes_applied}")
print(f"Table cells highlighted: {table_changes}")
print(f"Total changes marked: {changes_applied + table_changes}")

# ── Verification ───────────────────────────────────────────────────────────────

size = os.path.getsize(OUT_PATH)
print(f"\nFile size: {size/1024:.0f} KB")
print(f"{'OK — size reasonable' if size > 500_000 else 'WARNING — small file, check output'}")

doc_check = Document(OUT_PATH)
highlighted = 0
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
for p in doc_check.paragraphs:
    for run in p.runs:
        rPr = run._r.find(f'{{{NS}}}rPr')
        if rPr is not None and rPr.find(f'{{{NS}}}highlight') is not None:
            highlighted += 1
            break

print(f"Paragraphs with yellow highlight: {highlighted}")
print(
    f"{'OK — highlight applied correctly' if highlighted > 10 else 'WARNING — few highlights, check output'}"
)

print(
    "\nListo para subir a JADR como 'Revised Manuscript (with changes highlighted)'.\n"
    "Abrir en Word para verificación visual antes del envío."
)
